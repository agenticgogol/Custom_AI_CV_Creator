from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
import io
from typing import Optional, Dict, List
import re

# PDF generation imports
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False


class CVGenerator:
    """Generate downloadable CV files with enhanced formatting"""
    
    @staticmethod
    def create_docx_from_text(cv_text: str, changes_made: List[Dict] = None) -> bytes:
        """Convert CV text to professionally formatted DOCX"""
        doc = Document()
        
        # Set up document margins and styles
        CVGenerator._setup_document_styles(doc)
        
        # Parse CV text into sections
        sections = CVGenerator._parse_cv_sections(cv_text)
        
        # Add content with professional formatting
        for section_name, content in sections.items():
            if section_name.lower() == 'header':
                CVGenerator._add_header_section(doc, content)
            else:
                CVGenerator._add_section(doc, section_name, content, changes_made)
        
        # Add change summary if provided
        if changes_made:
            CVGenerator._add_changes_summary(doc, changes_made)
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    
    @staticmethod
    def create_pdf_from_text(cv_text: str, changes_made: List[Dict] = None) -> bytes:
        """Convert CV text to professionally formatted PDF"""
        if not HAS_REPORTLAB:
            raise ImportError("reportlab package required for PDF generation. Install with: pip install reportlab")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, 
                              rightMargin=0.75*inch, leftMargin=0.75*inch,
                              topMargin=0.75*inch, bottomMargin=0.75*inch)
        
        # Set up styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'],
                                   fontSize=16, spaceAfter=12, alignment=TA_CENTER,
                                   textColor=colors.darkblue)
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'],
                                     fontSize=14, spaceAfter=8, spaceBefore=12,
                                     textColor=colors.darkblue, 
                                     borderWidth=1, borderColor=colors.darkblue,
                                     borderPadding=2)
        body_style = ParagraphStyle('CustomBody', parent=styles['Normal'],
                                  fontSize=10, spaceAfter=4)
        bullet_style = ParagraphStyle('CustomBullet', parent=styles['Normal'],
                                    fontSize=10, leftIndent=20, bulletIndent=10,
                                    spaceAfter=3)
        
        # Parse CV content
        sections = CVGenerator._parse_cv_sections(cv_text)
        story = []
        
        # Add content
        for section_name, content in sections.items():
            if section_name.lower() == 'header':
                # Add contact information
                lines = [line.strip() for line in content.split('\n') if line.strip()]
                for line in lines[:3]:  # Name and contact info
                    if lines.index(line) == 0:  # Name
                        story.append(Paragraph(line, title_style))
                    else:  # Contact details
                        story.append(Paragraph(line, body_style))
                story.append(Spacer(1, 12))
            else:
                # Add section heading
                story.append(Paragraph(section_name.upper(), heading_style))
                
                # Add section content
                lines = [line.strip() for line in content.split('\n') if line.strip()]
                for line in lines:
                    if line.startswith(('•', '-', '*')):
                        story.append(Paragraph(line, bullet_style))
                    else:
                        story.append(Paragraph(line, body_style))
                story.append(Spacer(1, 8))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    @staticmethod
    def _setup_document_styles(doc):
        """Set up document margins and styles"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Create custom styles
        styles = doc.styles
        
        # Header style
        if 'CV Header' not in [style.name for style in styles]:
            header_style = styles.add_style('CV Header', WD_STYLE_TYPE.PARAGRAPH)
            header_style.font.name = 'Calibri'
            header_style.font.size = Pt(14)
            header_style.font.bold = True
            header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_style.paragraph_format.space_after = Pt(6)
        
        # Section heading style
        if 'CV Section' not in [style.name for style in styles]:
            section_style = styles.add_style('CV Section', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = 'Calibri'
            section_style.font.size = Pt(12)
            section_style.font.bold = True
            section_style.font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
    
    @staticmethod
    def _parse_cv_sections(cv_text: str) -> Dict[str, str]:
        """Parse CV text into sections"""
        sections = {}
        current_section = 'header'
        current_content = []
        
        lines = cv_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_content:  # Add empty line within content
                    current_content.append('')
                continue
            
            # Check if this is a section header
            if CVGenerator._is_section_header(line):
                # Save current section
                if current_content:
                    sections[current_section] = '\n'.join(current_content).strip()
                
                # Start new section
                current_section = line.replace(':', '').strip()
                current_content = []
            else:
                current_content.append(line)
        
        # Add final section
        if current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections
    
    @staticmethod
    def _is_section_header(line: str) -> bool:
        """Determine if a line is a section header"""
        headers = [
            'summary', 'objective', 'professional summary', 'experience', 
            'work experience', 'employment history', 'career history',
            'education', 'academic background', 'skills', 'technical skills',
            'core competencies', 'certifications', 'licenses', 'projects',
            'achievements', 'awards', 'publications', 'languages',
            'interests', 'references', 'contact', 'personal information',
            'volunteer experience', 'extracurricular activities'
        ]
        
        line_lower = line.lower().strip().replace(':', '')
        
        # Exact match with common headers
        if line_lower in headers:
            return True
        
        # Check if line is all caps (common for headers)
        if line.isupper() and len(line.split()) <= 4:
            return True
        
        # Check if line ends with colon
        if line.endswith(':') and len(line.split()) <= 4:
            return True
        
        return False
    
    @staticmethod
    def _add_header_section(doc, content):
        """Add formatted header section"""
        lines = [line.strip() for line in content.split('\n') if line.strip()]
        
        for i, line in enumerate(lines):
            if i == 0:  # Name
                p = doc.add_paragraph(line, style='CV Header')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:  # Contact info
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(3)
        
        # Add separator line
        doc.add_paragraph()
    
    @staticmethod
    def _add_section(doc, section_name, content, changes_made=None):
        """Add a formatted section"""
        # Add section heading
        heading = doc.add_paragraph(section_name.upper(), style='CV Section')
        
        # Add underline to section heading
        run = heading.runs[0]
        run.underline = True
        
        # Check if this section has changes
        section_changes = []
        if changes_made:
            section_changes = [
                change for change in changes_made 
                if change.get('section', '').lower() == section_name.lower()
            ]
        
        # Add content
        lines = [line.strip() for line in content.split('\n') if line.strip()]
        
        for line in lines:
            if line.startswith(('•', '-', '*')):
                # Bullet point
                p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                
                # Highlight if this was changed
                if CVGenerator._line_was_modified(line, section_changes):
                    for run in p.runs:
                        run.font.highlight_color = 3  # Yellow highlight
                        
            else:
                # Regular paragraph
                p = doc.add_paragraph(line)
                
                # Check for job titles, company names (bold formatting)
                if CVGenerator._is_job_title_or_company(line):
                    for run in p.runs:
                        run.font.bold = True
                
                # Highlight if this was changed
                if CVGenerator._line_was_modified(line, section_changes):
                    for run in p.runs:
                        run.font.highlight_color = 3  # Yellow highlight
    
    @staticmethod
    def _line_was_modified(line: str, section_changes: List[Dict]) -> bool:
        """Check if a line was modified based on changes"""
        for change in section_changes:
            if (change.get('change_type') in ['Added', 'Modified'] and 
                line in change.get('improved', '')):
                return True
        return False
    
    @staticmethod
    def _is_job_title_or_company(line: str) -> bool:
        """Check if line contains job title or company name"""
        # Simple heuristic: lines with job titles often have dates
        date_pattern = r'\b(19|20)\d{2}\b|\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)'
        has_date = bool(re.search(date_pattern, line, re.IGNORECASE))
        
        # Or contain position indicators
        position_indicators = [
            'engineer', 'manager', 'director', 'analyst', 'developer',
            'consultant', 'coordinator', 'specialist', 'lead', 'senior',
            'junior', 'associate', 'vice president', 'president', 'ceo',
            'cto', 'cfo', 'intern', 'trainee'
        ]
        
        has_position = any(indicator in line.lower() for indicator in position_indicators)
        
        return has_date or has_position
    
    @staticmethod
    def _add_changes_summary(doc, changes_made: List[Dict]):
        """Add a summary of changes made at the end"""
        if not changes_made:
            return
        
        doc.add_page_break()
        
        # Changes summary header
        heading = doc.add_paragraph('IMPROVEMENT SUMMARY', style='CV Section')
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("The following changes were made to optimize your resume for the target position:", 
                         style='Normal')
        doc.add_paragraph()
        
        # Group changes by section
        changes_by_section = {}
        for change in changes_made:
            section = change.get('section', 'General')
            if section not in changes_by_section:
                changes_by_section[section] = []
            changes_by_section[section].append(change)
        
        for section, section_changes in changes_by_section.items():
            # Section header
            section_para = doc.add_paragraph(f"{section.upper()}:", style='Heading 3')
            
            for change in section_changes:
                change_para = doc.add_paragraph(style='List Bullet')
                change_text = f"{change.get('change_type', 'Modified')}: {change.get('reason', 'Optimization')}"
                if change.get('addresses_gap'):
                    change_text += f" (Addresses: {change.get('addresses_gap')})"
                
                change_para.add_run(change_text)
    
    @staticmethod
    def create_text_file(cv_text: str, changes_made: List[Dict] = None) -> bytes:
        """Create a text file from CV text with optional changes summary"""
        content = cv_text
        
        if changes_made:
            content += "\n\n" + "="*50
            content += "\nIMPROVEMENT SUMMARY\n"
            content += "="*50 + "\n\n"
            content += "The following changes were made to optimize your resume:\n\n"
            
            changes_by_section = {}
            for change in changes_made:
                section = change.get('section', 'General')
                if section not in changes_by_section:
                    changes_by_section[section] = []
                changes_by_section[section].append(change)
            
            for section, section_changes in changes_by_section.items():
                content += f"{section.upper()}:\n"
                for change in section_changes:
                    content += f"  • {change.get('change_type', 'Modified')}: {change.get('reason', 'Optimization')}"
                    if change.get('addresses_gap'):
                        content += f" (Addresses: {change.get('addresses_gap')})"
                    content += "\n"
                content += "\n"
        
        return content.encode('utf-8')
    
    @staticmethod
    def validate_cv_text(cv_text: str) -> bool:
        """Validate that CV text contains essential elements"""
        cv_lower = cv_text.lower()
        
        # Check for essential sections
        required_elements = ['experience', 'education', 'skills']
        found_elements = sum(1 for element in required_elements if element in cv_lower)
        
        # Check for contact information
        has_contact = any(indicator in cv_lower for indicator in ['email', '@', 'phone'])
        
        return found_elements >= 2 and has_contact and len(cv_text.strip()) >= 200
    
    @staticmethod
    def create_comparison_document(original_cv: str, improved_cv: str, changes_made: List[Dict]) -> bytes:
        """Create a comparison document showing original vs improved"""
        doc = Document()
        CVGenerator._setup_document_styles(doc)
        
        # Title
        title = doc.add_paragraph('RESUME IMPROVEMENT COMPARISON', style='CV Header')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # Changes summary
        doc.add_paragraph('KEY IMPROVEMENTS MADE:', style='CV Section')
        
        for change in changes_made[:10]:  # Show top 10 changes
            bullet = doc.add_paragraph(style='List Bullet')
            bullet_text = f"{change.get('section', 'General')} - {change.get('change_type', 'Modified')}: {change.get('reason', 'Optimization')}"
            bullet.add_run(bullet_text)
        
        if len(changes_made) > 10:
            doc.add_paragraph(f"...and {len(changes_made) - 10} more improvements", style='Normal')
        
        doc.add_paragraph()
        
        # Side-by-side comparison would be complex in Word
        # Instead, show improved version with highlights
        doc.add_paragraph('IMPROVED RESUME:', style='CV Section')
        doc.add_paragraph('(Changed sections are highlighted)', style='Normal')
        
        # Add the improved CV with change tracking
        sections = CVGenerator._parse_cv_sections(improved_cv)
        for section_name, content in sections.items():
            if section_name.lower() != 'header':
                CVGenerator._add_section(doc, section_name, content, changes_made)
        
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()