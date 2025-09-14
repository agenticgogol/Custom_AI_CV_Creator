import mammoth
import io
from typing import Optional, Dict, List
from docx import Document
import logging

# Enhanced PDF parsing with multiple libraries
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

logger = logging.getLogger(__name__)


class DocumentParser:
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file using multiple methods"""
        text_results = []
        
        # Method 1: PyMuPDF (best for complex PDFs with tables/images)
        if HAS_PYMUPDF:
            try:
                text = DocumentParser._extract_with_pymupdf(file_content)
                if text and len(text.strip()) > 50:
                    text_results.append(("PyMuPDF", text))
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed: {e}")
        
        # Method 2: pdfplumber (good for tables)
        if HAS_PDFPLUMBER:
            try:
                text = DocumentParser._extract_with_pdfplumber(file_content)
                if text and len(text.strip()) > 50:
                    text_results.append(("pdfplumber", text))
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}")
        
        # Method 3: PyPDF2 (fallback)
        if HAS_PYPDF2:
            try:
                text = DocumentParser._extract_with_pypdf2(file_content)
                if text and len(text.strip()) > 50:
                    text_results.append(("PyPDF2", text))
            except Exception as e:
                logger.warning(f"PyPDF2 extraction failed: {e}")
        
        if not text_results:
            raise ValueError("Unable to extract text from PDF using any available method")
        
        # Return the longest extracted text (usually the best quality)
        best_result = max(text_results, key=lambda x: len(x[1]))
        logger.info(f"Best PDF extraction method: {best_result[0]}")
        
        return best_result[1].strip()
    
    @staticmethod
    def _extract_with_pymupdf(file_content: bytes) -> str:
        """Extract text using PyMuPDF (best for complex PDFs)"""
        doc = fitz.open(stream=file_content, filetype="pdf")
        text = ""
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            
            # Extract text blocks (preserves formatting better)
            blocks = page.get_text("dict")
            page_text = ""
            
            for block in blocks.get("blocks", []):
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        line_text = ""
                        for span in line.get("spans", []):
                            line_text += span.get("text", "")
                        if line_text.strip():
                            page_text += line_text + "\n"
                
                # Handle table data
                elif block.get("type") == 1:  # Image block - might contain table data
                    # Extract any text from image-based tables if possible
                    pass
            
            if page_text.strip():
                text += f"\n--- Page {page_num + 1} ---\n{page_text}"
        
        doc.close()
        return DocumentParser._clean_extracted_text(text)
    
    @staticmethod
    def _extract_with_pdfplumber(file_content: bytes) -> str:
        """Extract text using pdfplumber (excellent for tables)"""
        text = ""
        
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = ""
                
                # Extract regular text
                regular_text = page.extract_text()
                if regular_text:
                    page_text += regular_text + "\n"
                
                # Extract tables
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        page_text += "\n[TABLE]\n"
                        for row in table:
                            if row:
                                # Clean and join row data
                                clean_row = [str(cell).strip() if cell else "" for cell in row]
                                page_text += " | ".join(clean_row) + "\n"
                        page_text += "[/TABLE]\n"
                
                if page_text.strip():
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}"
        
        return DocumentParser._clean_extracted_text(text)
    
    @staticmethod
    def _extract_with_pypdf2(file_content: bytes) -> str:
        """Extract text using PyPDF2 (fallback method)"""
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text.strip():
                text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
        
        return DocumentParser._clean_extracted_text(text)
    
    @staticmethod
    def _clean_extracted_text(text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        
        # Remove empty lines and merge broken lines
        cleaned_lines = []
        for i, line in enumerate(lines):
            if not line:
                continue
            
            # If line ends with hyphen, likely a word break
            if line.endswith('-') and i < len(lines) - 1:
                next_line = lines[i + 1] if i + 1 < len(lines) else ""
                if next_line and not next_line[0].isupper():
                    line = line[:-1] + next_line
                    continue
            
            cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                text += "\n[TABLE]\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text += row_text + "\n"
                text += "[/TABLE]\n"
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_doc(file_content: bytes) -> str:
        """Extract text from DOC file using mammoth"""
        try:
            doc_file = io.BytesIO(file_content)
            result = mammoth.extract_raw_text(doc_file)
            return result.value.strip()
        except Exception as e:
            raise ValueError(f"Error parsing DOC: {str(e)}")
    
    @classmethod
    def parse_document(cls, file_content, filename: str) -> str:
        """Parse document based on file extension"""
        # Handle different types of file_content
        if hasattr(file_content, 'read'):
            # It's a file-like object (SpooledTemporaryFile)
            file_content.seek(0)  # Reset file pointer
            file_bytes = file_content.read()
        elif isinstance(file_content, bytes):
            # It's already bytes
            file_bytes = file_content
        else:
            # Try to convert to bytes
            try:
                file_bytes = bytes(file_content)
            except Exception:
                raise ValueError(f"Unsupported file content type: {type(file_content)}")
        
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return cls.extract_text_from_pdf(file_bytes)
        elif filename_lower.endswith('.docx'):
            return cls.extract_text_from_docx(file_bytes)
        elif filename_lower.endswith('.doc'):
            return cls.extract_text_from_doc(file_bytes)
        elif filename_lower.endswith('.txt'):
            return file_bytes.decode('utf-8')
        else:
            raise ValueError(f"Unsupported file format: {filename}")
    
    @staticmethod
    def validate_resume(text: str) -> bool:
        """Enhanced validation to check if text looks like a resume"""
        resume_keywords = [
            'experience', 'education', 'skills', 'work', 'employment',
            'university', 'college', 'degree', 'contact', 'email',
            'phone', 'address', 'objective', 'summary', 'projects',
            'responsibilities', 'achievements', 'professional', 'career'
        ]
        
        text_lower = text.lower()
        found_keywords = sum(1 for keyword in resume_keywords if keyword in text_lower)
        
        # Check for contact information patterns
        has_email = '@' in text and '.' in text
        has_phone = any(char.isdigit() for char in text) and (
            'phone' in text_lower or 'mobile' in text_lower or 'tel' in text_lower
        )
        
        # Check for date patterns (experience dates)
        import re
        date_pattern = r'\b(19|20)\d{2}\b|\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)'
        has_dates = bool(re.search(date_pattern, text_lower))
        
        # More stringent validation
        return (
            found_keywords >= 4 and 
            len(text.strip()) >= 200 and 
            (has_email or has_phone) and
            has_dates
        )
    @staticmethod
    def validate_job_description(text: str) -> bool:
        """Enhanced validation to check if text looks like a job description"""
        if not text or len(text.strip()) < 50:
            return False
            
        jd_keywords = [
            'requirements', 'responsibilities', 'qualifications', 'experience',
            'skills', 'required', 'preferred', 'must have', 'should have',
            'position', 'role', 'job', 'company', 'team', 'work', 'years',
            'candidate', 'looking for', 'seeking', 'we offer', 'benefits'
        ]
        
        text_lower = text.lower()
        found_keywords = sum(1 for keyword in jd_keywords if keyword in text_lower)
        
        # Check for job-specific patterns
        has_requirements = 'requirement' in text_lower or 'qualification' in text_lower
        has_responsibilities = 'responsibilit' in text_lower or 'duties' in text_lower
        has_experience_mention = 'year' in text_lower and 'experience' in text_lower
        
        return (
            found_keywords >= 4 and 
            len(text.strip()) >= 100 and
            (has_requirements or has_responsibilities) and
            has_experience_mention
        )


    @staticmethod
    def calculate_ats_compliance_score(text: str) -> Dict[str, any]:
        """Calculate ATS compliance using hybrid approach: rule-based + LLM analysis"""
        
        # First get rule-based score (quick, deterministic)
        rule_based_result = DocumentParser._calculate_rule_based_ats_score(text)
        
        # Only use LLM for borderline scores or when specifically requested
        # This saves API costs while providing enhanced analysis when needed
        use_llm = (
            rule_based_result["score"] >= 45 and rule_based_result["score"] <= 75
        ) or len(text) > 2000  # Use LLM for complex resumes or borderline scores
        
        if use_llm:
            # Then get LLM-based score if clients are available (more nuanced)
            llm_result = DocumentParser._calculate_llm_based_ats_score(text)
            
            # Combine both approaches
            if llm_result:
                # Weight: 30% rule-based, 70% LLM-based for final score
                final_score = int((rule_based_result["score"] * 0.3) + (llm_result["score"] * 0.7))
                
                # Combine feedback from both approaches
                combined_feedback = []
                combined_feedback.extend(llm_result["feedback"][:3])        # Top 3 LLM issues (prioritize)
                combined_feedback.extend(rule_based_result["feedback"][:2]) # Top 2 rule-based issues
                
                # Remove duplicates while preserving order
                seen = set()
                unique_feedback = []
                for item in combined_feedback:
                    if item.lower() not in seen:  # Case-insensitive duplicate check
                        unique_feedback.append(item)
                        seen.add(item.lower())
                
                return {
                    "score": final_score,
                    "grade": DocumentParser._get_ats_grade(final_score),
                    "feedback": unique_feedback[:5],  # Limit to top 5 recommendations
                    "rule_based_score": rule_based_result["score"],
                    "llm_score": llm_result["score"],
                    "scoring_method": "hybrid",
                    "sections_found": rule_based_result["sections_found"],
                    "contact_complete": rule_based_result["contact_complete"],
                    "has_quantified_achievements": rule_based_result["has_quantified_achievements"],
                    "action_verbs_count": rule_based_result["action_verbs_count"],
                    "detailed_analysis": llm_result.get("detailed_analysis", {})
                }
        
        # Fallback to rule-based only if LLM not used or failed
        rule_based_result["scoring_method"] = "rule_based"
        rule_based_result["rule_based_score"] = rule_based_result["score"]
        return rule_based_result
    
    @staticmethod 
    def _calculate_llm_based_ats_score(text: str) -> Dict[str, any]:
        """LLM-based ATS scoring for more nuanced analysis - optimized for cost"""
        
        try:
            from cv_agent.clients import get_client_analyzer, is_clients_initialized
            from langchain_core.messages import SystemMessage, HumanMessage
            
            if not is_clients_initialized():
                return None
            
            client = get_client_analyzer()
            
            # Truncate text to reduce API costs while keeping essential content
            truncated_text = text[:3000] + "..." if len(text) > 3000 else text
            
            prompt = f"""Analyze this resume for ATS compliance. Score 40-70 for typical resumes, 75+ only for exceptional ones.

RESUME:
{truncated_text}

Score based on:
1. Structure/Sections (25pts): Standard sections, clear headers
2. Contact Info (15pts): Proper email/phone formatting  
3. Format (20pts): ATS-friendly, no tables/complex formatting
4. Content (25pts): Action verbs, quantified achievements
5. Technical (15pts): Parsing compatibility

Return ONLY this JSON:
{{
    "score": 65,
    "confidence": "high",
    "feedback": ["Critical issue 1", "Issue 2", "Issue 3"],
    "detailed_analysis": {{
        "structure_score": 18,
        "contact_score": 12,
        "format_score": 15, 
        "content_score": 15,
        "technical_score": 5,
        "strengths": ["Strength 1"],
        "weaknesses": ["Weakness 1"],
        "keyword_density": "medium"
    }}
}}"""
            
            # Use proper LangChain message format
            response = client.invoke([
                SystemMessage(content="You are an ATS expert. Be realistic with scoring - most resumes score 40-70. Return only JSON."),
                HumanMessage(content=prompt)
            ])
            
            # Parse LLM response
            content = response.content.strip()
            if content.startswith("```json"):
                content = content[7:-3]
            elif content.startswith("```"):
                content = content[3:-3]
            
            import json
            result = json.loads(content)
            
            # Validate score is reasonable
            if result.get("score", 0) > 90:
                result["score"] = min(90, result["score"])  # Cap at 90 to prevent inflation
                
            return result
            
        except Exception as e:
            print(f"LLM ATS scoring failed: {e}")
            return None

    # Rest of the rule-based function stays the same...
    @staticmethod
    def _calculate_rule_based_ats_score(text: str) -> Dict[str, any]:
        """Rule-based ATS scoring - same implementation as before"""
        score = 0
        max_score = 100
        feedback = []
        text_lower = text.lower()
        
        # [Same implementation as in your document - keeping it unchanged]
        # ... (all the existing rule-based logic)
        
        # SECTION 1: Essential Structure (25 points)
        required_sections = {
            'contact': ['email', 'phone', '@', 'contact'],
            'experience': ['experience', 'work history', 'employment', 'career'],
            'education': ['education', 'degree', 'university', 'college', 'school'],
            'skills': ['skills', 'technical', 'proficiencies', 'competencies']
        }
        
        sections_found = 0
        missing_sections = []
        
        for section_name, keywords in required_sections.items():
            if any(keyword in text_lower for keyword in keywords):
                sections_found += 1
            else:
                missing_sections.append(section_name.title())
        
        if sections_found == 4:
            score += 25
        elif sections_found == 3:
            score += 15
            feedback.append(f"Missing critical section: {', '.join(missing_sections)}")
        elif sections_found == 2:
            score += 8
            feedback.append(f"Missing multiple sections: {', '.join(missing_sections)}")
        else:
            score += 0
            feedback.append(f"Missing essential sections: {', '.join(missing_sections)}")
        
        # SECTION 2: Contact Information (15 points)
        contact_score = 0
        import re
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        if re.search(email_pattern, text):
            contact_score += 8
        else:
            feedback.append("Missing valid email address format")
        
        phone_pattern = r'(\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'
        if re.search(phone_pattern, text):
            contact_score += 7
        else:
            feedback.append("Missing properly formatted phone number")
        
        score += contact_score
        
        # SECTION 3: Format and Structure (20 points)
        format_score = 0
        lines = text.split('\n')
        total_lines = len([line for line in lines if line.strip()])
        
        if total_lines < 15:
            feedback.append("Resume appears too short")
        elif total_lines > 100:
            feedback.append("Resume may be too long for ATS parsing")
            format_score -= 3
        else:
            format_score += 5
        
        # Bullet points check
        bullet_patterns = [r'•', r'-\s', r'\*\s', r'◦', r'▪']
        bullet_count = sum(len(re.findall(pattern, text)) for pattern in bullet_patterns)
        
        if bullet_count >= 8:
            format_score += 5
        elif bullet_count >= 4:
            format_score += 3
            feedback.append("Could use more bullet points")
        else:
            feedback.append("Insufficient bullet points")
        
        # Formatting problems
        if '[TABLE]' in text or text.count('|') > 5:
            feedback.append("Contains tables - may cause ATS parsing issues")
            format_score -= 2
        
        if re.search(r'^[A-Z\s]{3,20}$', text, re.MULTILINE):
            format_score += 5
        else:
            feedback.append("Missing clear section headers")
        
        format_score = max(0, format_score)
        score += min(20, format_score)
        
        # SECTION 4: Content Quality (25 points)
        content_score = 0
        
        action_verbs = [
            'achieved', 'managed', 'led', 'developed', 'implemented', 'created',
            'improved', 'increased', 'reduced', 'coordinated', 'designed',
            'built', 'executed', 'delivered', 'optimized'
        ]
        
        found_verbs = [verb for verb in action_verbs if verb in text_lower]
        unique_verbs = len(set(found_verbs))
        
        if unique_verbs >= 8:
            content_score += 8
        elif unique_verbs >= 5:
            content_score += 5
            feedback.append("Could use more varied action verbs")
        else:
            feedback.append("Insufficient action verbs")
        
        # Quantified achievements
        number_patterns = [
            r'\d{1,3}%', r'\$[\d,]+', r'\d+\+?\s*(years?|months?)',
            r'\d{1,3}[kmb]?\s*(users?|customers?|people)',
            r'(increased|improved|reduced|saved).*?\d+',
            r'\d+\s*(projects?|teams?|reports?)'
        ]
        
        quantified_count = 0
        for pattern in number_patterns:
            quantified_count += len(re.findall(pattern, text, re.IGNORECASE))
        
        if quantified_count >= 6:
            content_score += 10
        elif quantified_count >= 3:
            content_score += 6
            feedback.append("Add more quantified achievements")
        else:
            feedback.append("Lacks quantified achievements")
        
        if len(text.split()) >= 200:
            content_score += 4
        else:
            feedback.append("Resume content appears too brief")
        
        unprofessional_words = ['stuff', 'things', 'lots', 'many', 'various', 'etc']
        unprofessional_found = sum(1 for word in unprofessional_words if word in text_lower)
        
        if unprofessional_found > 2:
            content_score -= 3
            feedback.append("Contains vague language")
        else:
            content_score += 3
        
        score += min(25, content_score)
        
        # SECTION 5: Technical Compatibility (15 points)
        technical_score = 15
        
        if len(text) < 500:
            technical_score -= 5
            feedback.append("Resume appears too short")
        
        problematic_chars = ['â€™', 'â€œ', 'â€', '�', 'Â']
        if any(char in text for char in problematic_chars):
            technical_score -= 3
            feedback.append("Contains character encoding issues")
        
        if text.count('\n\n') > text.count('\n') * 0.3:
            technical_score -= 2
            feedback.append("Excessive spacing may confuse ATS")
        
        score += max(0, technical_score)
        
        final_score = min(100, int(score))
        
        return {
            "score": final_score,
            "grade": DocumentParser._get_ats_grade(final_score),
            "feedback": feedback,
            "sections_found": sections_found,
            "contact_complete": contact_score >= 12,
            "has_quantified_achievements": quantified_count >= 3,
            "action_verbs_count": unique_verbs,
            "quantified_achievements_count": quantified_count,
            "technical_issues": 15 - max(0, technical_score),
            "max_possible_score": max_score
        }
    
    @staticmethod
    def _get_ats_grade(score: int) -> str:
        """Convert numeric score to grade"""
        if score >= 85:
            return "Excellent"
        elif score >= 75:
            return "Good"
        elif score >= 60:
            return "Fair"
        elif score >= 40:
            return "Poor"
        else:
            return "Very Poor"